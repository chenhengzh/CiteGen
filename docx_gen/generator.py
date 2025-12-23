import json
import os
import logging
import docx
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement

from utils import get_filename, list_data_in_directory, are_strings_almost_matching
from .downloader import get_pdf
from config import PAPER_LIST_DIR


def display_cit(cit):
    logging.info(
        "+++===================================================================================================+++"
    )
    logging.info(f"index: {cit.get('index', 'N/A')}")
    logging.info(f"title: {cit.get('title', 'N/A')}")
    logging.info(f"filename: {cit.get('filename', 'N/A')}")
    logging.info(f"info: {cit.get('info', 'N/A')}")
    logging.info(f"abstract: {cit.get('abstract', 'N/A')}")
    if not cit.get("PDF"):
        logging.info("no PDF resource.")
    else:
        logging.info(f"PDF: {cit['PDF']}")
    logging.info(f"paper_link: {cit.get('link', 'N/A')}")


# word format related functions
def add_hyperlink(paragraph, text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(
        url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True
    )

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement("w:hyperlink")
    hyperlink.set(
        docx.oxml.shared.qn("r:id"),
        r_id,
    )

    # Create a new run object (a wrapper over a 'w:r' element)
    new_run = docx.text.run.Run(docx.oxml.shared.OxmlElement("w:r"), paragraph)
    new_run.text = text

    # Set the run's style to the builtin hyperlink style, defining it if necessary
    new_run.style = get_or_create_hyperlink_style(part.document)

    # Join all the xml elements together
    hyperlink.append(new_run._element)
    paragraph._p.append(hyperlink)
    return hyperlink


def get_or_create_hyperlink_style(d):
    """If this document had no hyperlinks so far, the builtin
    Hyperlink style will likely be missing and we need to add it.
    """

    if "Hyperlink" not in d.styles:
        if "Default Character Font" not in d.styles:
            ds = d.styles.add_style(
                "Default Character Font", docx.enum.style.WD_STYLE_TYPE.CHARACTER, True
            )
            ds.element.set(docx.oxml.shared.qn("w:default"), "1")
            ds.priority = 1
            ds.hidden = True
            ds.unhide_when_used = True
            del ds
        hs = d.styles.add_style(
            "Hyperlink", docx.enum.style.WD_STYLE_TYPE.CHARACTER, True
        )
        hs.base_style = d.styles["Default Character Font"]
        hs.unhide_when_used = True
        hs.font.color.rgb = docx.shared.RGBColor(0, 0, 255)
        hs.font.underline = True
        hs.font.size = Pt(13)
        hs.font.name = "Arial"
        del hs

    return "Hyperlink"


def get_locallink(cit, pdf_list):
    # paper_file=get_filename(paper.title)
    for pdf in pdf_list:
        ismatch = are_strings_almost_matching(cit["filename"], pdf[:-4], threshold=90)
        if ismatch:
            return pdf
    return ""


def input_docx(cit, doc_pth, is_pdf, pdf_list=None):
    if pdf_list is None:
        pdf_list = []

    logging.info("+======writing item======+")

    doc = Document(doc_pth)
    is_written = bool(len(doc.paragraphs))
    para = doc.add_paragraph()
    # set the space before the paragraph
    if is_written:
        para.paragraph_format.space_before = Pt(16)
    # set the first line indent
    para.paragraph_format.first_line_indent = Pt(0)
    # set the hanging indent
    para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    line1_text = cit["title"] + "\n"
    if is_pdf:
        line1_link = cit["filename"] + ".pdf"
    else:
        pdf_link = get_locallink(cit, pdf_list)
        if pdf_link:
            line1_link = pdf_link
        else:
            line0 = "[PDF not downloaded]\n"
            run0 = para.add_run(line0)
            run0.font.name = "Arial"
            run0.font.size = Pt(12)
            run0.font.color.rgb = RGBColor(200, 0, 0)  # red
            line1_link = cit["link"]

    add_hyperlink(para, line1_text, line1_link)

    line2 = cit["info"] + "\n"
    run2 = para.add_run(line2)
    run2.font.name = "Arial"
    run2.font.size = Pt(10)
    run2.font.color.rgb = RGBColor(0, 102, 33)  # green

    line3 = cit["abstract"]
    run3 = para.add_run(line3)
    run3.font.name = "Arial"
    run3.font.size = Pt(10)
    run3.font.color.rgb = RGBColor(34, 34, 34)  # black

    doc.save(doc_pth)

    logging.info("+======item done======+")


def docx_worker(paper_title, get_pdf_flag=True):
    print(
        f"***++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++***"
    )
    print(f"Start writing the docx document of the paper: [{paper_title}]")
    print(
        f"***++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++***\n"
    )

    dir_name = get_filename(paper_title)
    json_path = os.path.join(PAPER_LIST_DIR, dir_name, "citation_info.json")

    if os.path.exists(json_path):
        with open(json_path, "r") as file:
            cit_list = json.load(file)
    else:
        cit_list = []

    isPDF = 0

    logging.info("\n\n\n")
    logging.info(
        f"\n***+++++++++++++++++++++++++++++writing the docx of Paper: [{dir_name}]+++++++++++++++++++++++++++++***\n"
    )
    if not cit_list:
        logging.info(f"Paper: [{dir_name}] has no citation")
        print(
            f"***++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++***"
        )
        print(f"Paper: [{dir_name}] has no citation")
        print(
            f"***++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++***"
        )
        print()
        return

    if get_pdf_flag:
        doc_pth = os.path.join(PAPER_LIST_DIR, dir_name, f"(temp) {dir_name}.docx")
    else:
        doc_pth = os.path.join(PAPER_LIST_DIR, dir_name, f"{dir_name}.docx")

    # Always overwrite for fresh generation
    doc = Document()
    doc.save(doc_pth)

    pdf_files = []
    if not get_pdf_flag:
        dir_path = os.path.dirname(doc_pth)
        # Record downloaded pdf files for relative address modification
        if os.path.exists(dir_path):
            pdf_files = [file for file in os.listdir(dir_path) if file.endswith(".pdf")]

    for cit in cit_list:
        display_cit(cit)
        pdf_pth = os.path.join(PAPER_LIST_DIR, dir_name, f"{cit['filename']}.pdf")

        if get_pdf_flag:
            # Check if PDF already exists before downloading
            if os.path.exists(pdf_pth) and os.path.getsize(pdf_pth) > 0:
                logging.info(f"PDF already exists at {pdf_pth}, skipping download.")
                isPDF = True
            else:
                isPDF = get_pdf(cit, pdf_pth)

            input_docx(cit, doc_pth, isPDF)
        else:
            input_docx(cit, doc_pth, False, pdf_list=pdf_files)

        logging.info(
            "+++===================================================================================================+++\n"
        )

    print(
        f"***++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++***"
    )
    print(
        f"The docx document of the paper: [{paper_title}] has been written successfully."
    )
    print(
        f"***++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++***\n"
    )


def generate_all_docx(paper_ls, get_pdf_flag=True):
    print()
    print(f"The {str(len(paper_ls))} docx documents to be written:")
    print(paper_ls)
    print(
        "+++===================================================================================================+++"
    )
    print()

    logging.info("\n\n\n")
    logging.info(
        f"#####***++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++***#####"
    )
    logging.info(f"The following is a new process")
    logging.info(
        f"#####***++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++***#####"
    )
    logging.info("\n\n\n")

    for paper in paper_ls:
        docx_worker(paper, get_pdf_flag=get_pdf_flag)
    print("All docx documents have been written successfully.")

    logging.info("\n\n\n")
    logging.info(
        f"#####***++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++***#####"
    )
    logging.info(f"All docx documents have been written successfully.")
    logging.info(
        f"#####***++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++***#####"
    )
    logging.info("\n\n\n")
