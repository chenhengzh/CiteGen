import json
import os
import logging
import docx
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement

from utils import get_filename, are_strings_almost_matching
from config import PAPER_LIST_DIR, CITATION_HIGHLIGHT_THRESHOLD

# Color definitions
COLOR_BLUE = RGBColor(0, 0, 255)
COLOR_ORANGE = RGBColor(255, 128, 0)
COLOR_RED = RGBColor(255, 0, 0)
COLOR_DARK_RED = RGBColor(200, 0, 0)
COLOR_GREEN = RGBColor(0, 102, 33)
COLOR_BLACK = RGBColor(34, 34, 34)
COLOR_GREY = RGBColor(50, 50, 50)
COLOR_DARK_BLUE = RGBColor(0, 0, 139)


def display_cit(cit):
    logging.info(
        "+++===================================================================================================+++"
    )
    logging.info(f"index: {cit.get('index', 'N/A')}")
    logging.info(f"title: {cit.get('title', 'N/A')}")
    logging.info(f"filename: {cit.get('filename', 'N/A')}")
    logging.info(f"info: {cit.get('info', 'N/A')}")
    logging.info(f"abstract: {cit.get('abstract', 'N/A')}")
    if not cit.get("PDF"):
        logging.info("no PDF resource.")
    else:
        logging.info(f"PDF: {cit['PDF']}")
    logging.info(f"paper_link: {cit.get('link', 'N/A')}")


# word format related functions
def add_hyperlink(
    paragraph, text, url, color=COLOR_BLUE, font_size=None, font_name="Arial"
):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(
        url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True
    )

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement("w:hyperlink")
    hyperlink.set(
        docx.oxml.shared.qn("r:id"),
        r_id,
    )

    # Create a new run object (a wrapper over a 'w:r' element)
    new_run = docx.text.run.Run(docx.oxml.shared.OxmlElement("w:r"), paragraph)
    new_run.text = text

    # Set the run's style to the builtin hyperlink style, defining it if necessary
    new_run.style = get_or_create_hyperlink_style(part.document)
    new_run.font.color.rgb = color
    if font_size:
        new_run.font.size = font_size
    if font_name:
        new_run.font.name = font_name

    # Join all the xml elements together
    hyperlink.append(new_run._element)
    paragraph._p.append(hyperlink)
    return hyperlink


def get_or_create_hyperlink_style(d):
    """If this document had no hyperlinks so far, the builtin
    Hyperlink style will likely be missing and we need to add it.
    """

    if "Hyperlink" not in d.styles:
        if "Default Character Font" not in d.styles:
            ds = d.styles.add_style(
                "Default Character Font", docx.enum.style.WD_STYLE_TYPE.CHARACTER, True
            )
            ds.element.set(docx.oxml.shared.qn("w:default"), "1")
            ds.priority = 1
            ds.hidden = True
            ds.unhide_when_used = True
            del ds
        hs = d.styles.add_style(
            "Hyperlink", docx.enum.style.WD_STYLE_TYPE.CHARACTER, True
        )
        hs.base_style = d.styles["Default Character Font"]
        hs.unhide_when_used = True
        hs.font.color.rgb = COLOR_BLUE
        hs.font.underline = True
        # hs.font.size = Pt(13)
        hs.font.name = "Arial"
        del hs

    return "Hyperlink"


def get_locallink(cit, pdf_list):
    # paper_file=get_filename(paper.title)
    for pdf in pdf_list:
        ismatch = are_strings_almost_matching(cit["filename"], pdf[:-4], threshold=90)
        if ismatch:
            return pdf
    return ""


def input_docx(cit, doc_pth, pdf_filename=None):
    """
    Writes a citation into the docx.
    pdf_filename: If provided, links to this local file. If None, links to web URL and marks as not downloaded.
    """

    logging.info("+======writing item======+")

    doc = Document(doc_pth)
    doc.add_paragraph()
    is_written = bool(len(doc.paragraphs))
    para = doc.add_paragraph()
    # set the space before the paragraph
    if is_written:
        para.paragraph_format.space_before = Pt(16)
    # set the first line indent
    para.paragraph_format.first_line_indent = Pt(0)
    # set the hanging indent
    para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    line1_text = cit["title"] + "\n"

    if pdf_filename:
        line1_link = pdf_filename
    else:
        line0 = "[PDF not downloaded]\n"
        run0 = para.add_run(line0)
        run0.font.name = "Arial"
        run0.font.size = Pt(12)
        run0.font.color.rgb = COLOR_DARK_RED  # red
        line1_link = cit["link"]

    add_hyperlink(para, line1_text, line1_link, font_size=Pt(12))

    line2 = cit["info"] + "\n"
    run2 = para.add_run(line2)
    run2.font.name = "Arial"
    run2.font.size = Pt(10)
    run2.font.color.rgb = COLOR_GREEN  # green

    line3 = cit["abstract"]
    run3 = para.add_run(line3)
    run3.font.name = "Arial"
    run3.font.size = Pt(10)
    run3.font.color.rgb = COLOR_BLACK  # black

    # 1. Add Author Citation Info
    authors = cit.get("authors", [])
    author_status = cit.get("author_status", "")

    if authors or author_status == "omitted":
        para_auth = doc.add_paragraph()
        para_auth.paragraph_format.first_line_indent = Pt(0)
        para_auth.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        for author in authors:
            name = author.get("name", "Unknown")
            citations = author.get("citations", 0)
            link = author.get("link", "")

            # Format: "name: citations"
            text_str = f"{name}: {citations}\n"

            is_highlight = False
            try:
                if int(citations) > CITATION_HIGHLIGHT_THRESHOLD:
                    is_highlight = True
            except (ValueError, TypeError):
                pass

            if link:
                add_hyperlink(
                    para_auth,
                    text_str,
                    link,
                    color=COLOR_ORANGE if is_highlight else COLOR_BLACK,
                    font_size=Pt(10),
                )
            else:
                run_auth = para_auth.add_run(text_str)
                run_auth.font.name = "Arial"
                run_auth.font.size = Pt(10)
                run_auth.font.color.rgb = COLOR_ORANGE if is_highlight else COLOR_BLACK

        if author_status == "omitted":
            run_omit = para_auth.add_run("[Author Info Incomplete]\n")
            run_omit.font.name = "Arial"
            run_omit.font.size = Pt(10)
            run_omit.italic = True
            run_omit.font.color.rgb = COLOR_DARK_RED

    # 2. Add Positive Citations Analysis
    paper_dir = os.path.dirname(doc_pth)
    analysis_filename = cit.get("filename", "")
    if analysis_filename:
        analysis_path = os.path.join(
            paper_dir, "comment_analysis", f"{analysis_filename}.json"
        )

        if os.path.exists(analysis_path):
            try:
                with open(analysis_path, "r") as f:
                    analysis_data = json.load(f)

                positive_cits = [
                    c
                    for c in analysis_data.get("Citations", [])
                    if c.get("Positive") is True
                ]

                if positive_cits:
                    # Empty line
                    # doc.add_paragraph()

                    # Header for Positive Citations (Optional, but good for structure)
                    # para_header = doc.add_paragraph()
                    # run_header = para_header.add_run("Positive Citations:")
                    # run_header.bold = True

                    for pc in positive_cits:
                        para_pc = doc.add_paragraph()
                        para_pc.paragraph_format.first_line_indent = Pt(0)

                        text = pc.get("Text", "").strip()
                        analysis = pc.get("Analysis", "").strip()

                        run_text = para_pc.add_run(f'"{text}"\n')
                        run_text.font.name = "Arial"
                        run_text.font.size = Pt(10)
                        run_text.italic = True
                        run_text.font.color.rgb = COLOR_GREY

                        if analysis:
                            run_analysis = para_pc.add_run(f"Analysis: {analysis}")
                            run_analysis.font.name = "Microsoft YaHei"
                            run_analysis._element.rPr.rFonts.set(
                                docx.oxml.shared.qn("w:eastAsia"), "Microsoft YaHei"
                            )
                            run_analysis.font.size = Pt(10)
                            run_analysis.font.color.rgb = (
                                COLOR_DARK_BLUE  # Dark Blue for analysis
                            )

            except Exception as e:
                logging.error(f"Error reading analysis file {analysis_path}: {e}")

    doc.save(doc_pth)

    logging.info("+======item done======+")


def report_worker(paper_title):
    print(
        f"***++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++***"
    )
    print(f"Start generating report for paper: [{paper_title}]")
    print(
        f"***++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++***\n"
    )

    dir_name = get_filename(paper_title)
    paper_dir = os.path.join(PAPER_LIST_DIR, dir_name)
    json_path = os.path.join(paper_dir, "citation_info.json")

    if os.path.exists(json_path):
        with open(json_path, "r") as file:
            cit_list = json.load(file)
    else:
        cit_list = []

    logging.info("\n\n\n")
    logging.info(
        f"\n***+++++++++++++++++++++++++++++writing the docx of Paper: [{dir_name}]+++++++++++++++++++++++++++++***\n"
    )
    if not cit_list:
        logging.info(f"Paper: [{dir_name}] has no citation")
        print(f"Paper: [{dir_name}] has no citation")
        return

    doc_pth = os.path.join(paper_dir, f"{dir_name}.docx")

    # Always overwrite for fresh generation
    doc = Document()
    doc.save(doc_pth)

    # List all PDF files in the directory for fuzzy matching fallback
    pdf_files = []
    if os.path.exists(paper_dir):
        pdf_files = [file for file in os.listdir(paper_dir) if file.endswith(".pdf")]

    for cit in cit_list:
        display_cit(cit)

        # Determine PDF link
        pdf_filename = None

        # 1. Check exact match
        exact_pdf_name = f"{cit['filename']}.pdf"
        if os.path.exists(os.path.join(paper_dir, exact_pdf_name)):
            pdf_filename = exact_pdf_name
        else:
            # 2. Check fuzzy match
            fuzzy_match = get_locallink(cit, pdf_files)
            if fuzzy_match:
                pdf_filename = fuzzy_match

        input_docx(cit, doc_pth, pdf_filename)

        logging.info(
            "+++===================================================================================================+++\n"
        )

    print(
        f"***++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++***"
    )
    print(
        f"The docx document of the paper: [{paper_title}] has been written successfully."
    )
    print(
        f"***++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++***\n"
    )


def generate_all_reports(paper_ls):
    print()
    print(f"The {str(len(paper_ls))} reports to be generated:")
    print(paper_ls)
    print(
        "+++===================================================================================================+++"
    )
    print()

    logging.info("\n\n\n")
    logging.info(
        f"#####***++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++***#####"
    )
    logging.info(f"The following is a new report generation process")
    logging.info(
        f"#####***++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++***#####"
    )
    logging.info("\n\n\n")

    for paper in paper_ls:
        report_worker(paper)
    print("All reports have been generated successfully.")

    logging.info("\n\n\n")
    logging.info(
        f"#####***++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++***#####"
    )
    logging.info(f"All reports have been generated successfully.")
    logging.info(
        f"#####***++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++***#####"
    )
    logging.info("\n\n\n")
